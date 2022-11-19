#!/usr/bin/env python
# coding: utf-8

# In[1]:


from docx import Document


# In[2]:


def read_qas(file):
    qas_doc = Document(file)
    questions = []
    for para in qas_doc.paragraphs:
        question = {}
        k, v = para.text.split(":", 1)
        question[k.strip()] = v.strip()
        questions.append(question)
    return questions


# In[3]:


def print_result(name, ans):
    output = Document()

    output.add_heading(f"{name} - Review Result", 0)

    output.add_heading("Part 1. Basic Information",1)

    for i in range(0,7):
        q = "Q" + str(i+1) + ". " + ans[i]["query"]
        a = ans[i]["answers"][0].answer
        output.add_heading(q,2)
        output.add_heading(a,2)
        output.add_heading(" ",2)
        
    
    q8 = "Q8. " + ans[7]["query"]
    a8 = ans[7]["answers"][0].answer
    output.add_heading(q8,2)
    output.add_heading(a8,2)
    if "China" not in a8 and "New York" not in a8:
        output.add_heading("Attention: the governing law may not acceptable.",2)
        
    output.add_heading("Part 2. Important clauses for review",1)
    for i in range(8,13):
        q = "Q" + str(i+1) + ". " + ans[i]["query"]
        a = ans[i]["answers"][0].answer
        output.add_heading(q,2)
        output.add_heading(a,2)
        output.add_heading(" ",2)
    
    output_dir = f"./Review_Results/{name}_Result.docx"
    output.save(output_dir)


# In[ ]:




