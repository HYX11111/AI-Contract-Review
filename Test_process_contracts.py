#!/usr/bin/env python
# coding: utf-8

# In[1]:


from docx import Document
import torch
import transformers 
from haystack.document_stores import ElasticsearchDocumentStore
from haystack.nodes import BM25Retriever, TfidfRetriever, FilterRetriever, DensePassageRetriever
from haystack.nodes import FARMReader
from haystack.pipelines import ExtractiveQAPipeline
from haystack.utils import print_answers
import spacy
nlp = spacy.load("en_core_web_sm")
import re


# In[2]:


import import_ipynb
import Test_models as test_models
from Load_and_Print import read_qas, print_result


# ##### After testing differrent models for retrivers and readers, model 3 has been chosen.
# ##### Based on model 3, test whether to to tokenize the text of the contract.

# In[3]:


def process_contract1(file_dir):
    contract = Document(file_dir)
    
    contract_main = ""
    
    for para in contract.paragraphs:
        if check_para(para.text):
            contract_main = contract_main + para.text
    
    return contract_main


# In[4]:


def process_contract2(file_dir):
    contract = Document(file_dir)
    
    contract_main = ""
    
    for para in contract.paragraphs:
        if check_para(para.text):
            contract_main = contract_main + para.text
    
    return tokenize(contract_main)


# In[5]:


def check_para(p):
    if p.lower().startswith('ex') and not p == '':
        return False
    elif len(p)<=5:
        return False
    elif 'WITNESS ' in p or 'in witness' in p or 'signature' in p or 'signatures' in p:
        return False
    else:
        return True


# In[6]:


def tokenize(text):
    tokenized_texts = []
    text = re.sub("[^a-zA-Z0-9]", " ", text)
        
    tokens = nlp(text)
    for token in tokens:
        if not token.is_stop and not token.is_punct:
            tokenized_texts.append(token.lemma_.lower())
    return " ".join(tokenized_texts)
    


# In[7]:


qas = read_qas("./Questions.docx")


# In[8]:


def test1():
    qas = read_qas('Questions.docx')
    for i in range(0,3):
        input_dir = f"./Agreements/Test Agreement {i+1}.docx"
        contract1 = process_contract1(input_dir)
        print(f"Contract{i+1}")
        doc_store = test_models.set_doc_store(contract1, 1)
        ans = test_models.test_model3(qas, contract1, doc_store)
        doc_store.delete_index(str(1))


# In[9]:


def test2():
    qas = read_qas('Questions.docx')
    for i in range(0,3):
        input_dir = f"./Agreements/Test Agreement {i+1}.docx"
        contract2 = process_contract2(input_dir)
        print(f"Contract{i+1}")
        doc_store = test_models.set_doc_store(contract2, 2)
        ans = test_models.test_model3(qas, contract2, doc_store)
        doc_store.delete_index(str(2))


# In[10]:


if __name__ == '__main__':
    test1()


# In[11]:


if __name__ == '__main__':
    test2()


# In[ ]:





# In[ ]:




